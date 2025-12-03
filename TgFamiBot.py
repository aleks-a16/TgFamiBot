from telegram import Update, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler
from docx import Document
import io, os

# СОСТОЯНИЯ СРАНОГО ДИАЛОГА
s_template, f_data = range(3)

# ХРАНИЛИЩЕ
data_store = {}

# ШАБЛОНЫЫЫЫЫЫ
TEMPLATES = {
    "CK_FZ": {
        "name": "Служебная записка на внос/вынос ФЗ в ЦК",
        "file": "CK_FZ.docx",
        "fields": [
            ("in_date", "Дата вноса: (дд месяца гггг)"),
            ("out_date", "Дата выноса:"),
            ("event", "Название мероприятия:")
        ]
    },
    "Predmety": {
        "name": "Служебка на внос/вынос вещей",
        "file": "Predmety.docx",
        "fields": [
            ("event", "Название мероприятия:"),
            ("in_date", "Дата вноса:"),
            ("place", "Место вноса в Р.п:"),
            ("thing", "Предметы через запятую:")]
    },
    "FZ_OK": {
        "name": "Разрешение на ФЗ",
        "file": "FZ_OK.docx",
        "fields": [
            ("st_date", "Дата начала:"),
            ("fin_date", "Дата конца"),
            ("event", "Название мероприятия")
        ]
    }
}

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # Создаем клавиатуру с вариантами шаблонов
    keyboard = [
        [KeyboardButton(TEMPLATES["CK_FZ"]["name"])],
        [KeyboardButton(TEMPLATES["Predmety"]["name"])],
        [KeyboardButton(TEMPLATES["FZ_OK"]["name"])]
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    
    await update.message.reply_text(
        "🤖 ДОБРО ПОЖАЛОВАТЬ В СИСТЕМУ СОЗДАНИЯ ДОКУМЕНТОВ!\n\n"
        "Выберите тип служебной записки:",
        reply_markup=reply_markup
    )
    
    return s_template

async def handle_template_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_text = update.message.text
    
    # Определяем какой шаблон выбрал пользователь
    selected_template = None
    for key, template in TEMPLATES.items():
        if template["name"] == user_text:
            selected_template = key
            break
    
    if not selected_template:
        await update.message.reply_text("❌ Пожалуйста, выберите шаблон из предложенных вариантов.")
        return s_template
    
    # Сохраняем выбранный шаблон в данных пользователя
    user_id = update.message.from_user.id
    data_store[user_id] = {
        'template': selected_template,
        'current_field': 0,
        'data': {}
    }
    
    # Получаем первый вопрос для заполнения
    template_info = TEMPLATES[selected_template]
    first_field = template_info["fields"][0]
    
    await update.message.reply_text(
        f"✅ Выбран: {template_info['name']}\n\n"
        f"Заполните данные:\n"
        f"{first_field[1]}"
    )
    
    return f_data

def create_filled_document(template_file, data):
    """
    Заполняет Word шаблон данными пользователя
    """
    try:
        # ПРОВЕРЯЕМ СУЩЕСТВОВАНИЕ ФАЙЛА
        if not os.path.exists(template_file):
            # Пытаемся найти файл в разных местах
            possible_paths = [
                template_file,
                f"templates/{template_file}"]
            
            found_path = None
            for path in possible_paths:
                if os.path.exists(path):
                    found_path = path
                    break
            
            if not found_path:
                raise FileNotFoundError(f"Файл шаблона '{template_file}' не найден. Проверьте наличие файлов: {', '.join(possible_paths)}")
            
            template_file = found_path
        
        # Открываем шаблон
        doc = Document(template_file)
        
        # Функция для замены текста в параграфах
        def replace_in_paragraph(paragraph, data):
            for run in paragraph.runs:
                original_text = run.text
                new_text = original_text
                
                for key, value in data.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, str(value))
                
                if new_text != original_text:
                    run.text = new_text
        
        # Заменяем в обычных параграфах
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, data)
        
        # Заменяем в таблицах (если есть)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, data)
        
        # Сохраняем в буфер
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except FileNotFoundError as e:
        raise Exception(f"Файл шаблона не найден: {e}")
    except Exception as e:
        raise Exception(f"Ошибка при заполнении шаблона: {e}")
    
async def handle_data_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    # Проверяем, есть ли данные пользователя
    if user_id not in data_store:
        await update.message.reply_text("❌ Сессия устарела. Начните заново с /start")
        return ConversationHandler.END
    
    user_data = data_store[user_id]
    template_key = user_data['template']
    template_info = TEMPLATES[template_key]
    fields = template_info["fields"]
    current_index = user_data['current_field']
    
    # Сохраняем введенные данные
    current_field_name = fields[current_index][0]
    user_data['data'][current_field_name] = update.message.text
    
    # Переходим к следующему полю или завершаем
    next_index = current_index + 1
    
    if next_index < len(fields):
        # Есть еще поля для заполнения
        user_data['current_field'] = next_index
        next_field = fields[next_index]
        
        await update.message.reply_text(
            f"✅ Сохранено!\n\n"
            f"Следующий вопрос:\n"
            f"{next_field[1]}"
        )
        
        return f_data
    else:
        # Все поля заполнены - создаем документ
        await update.message.reply_text("📄 Создаю документ...")
        
        # Создаем документ
        try:
            document_buffer = create_filled_document(template_info["file"], user_data['data'])
            
            # Отправляем документ пользователю
            await update.message.reply_document(
                document=document_buffer,
                filename=f"{template_info['name']}.docx",
                caption="✅ Ваш документ готов!\n\n"
                       "Для создания нового документа используйте /start"
            )
            
        except Exception as e:
            await update.message.reply_text(f"❌ Ошибка при создании документа: {e}")
        
        # Очищаем данные пользователя
        del data_store[user_id]
        
        return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Отмена операции
    """
    user_id = update.message.from_user.id
    if user_id in data_store:
        del data_store[user_id]
    
    await update.message.reply_text(
        "❌ Операция отменена.\n"
        "Для начала заново используйте /start",
        reply_markup=None
    )
    
    return ConversationHandler.END

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """
    Справка по использованию бота
    """
    await update.message.reply_text(
        "🤖 БОТ ДЛЯ СОЗДАНИЯ СЛУЖЕБНЫХ ЗАПИСОК\n\n"
        "Доступные команды:\n"
        "/start - начать создание документа\n"
        "/help - эта справка\n\n"
        "Как использовать:\n"
        "1. Выберите тип документа\n"
        "2. Последовательно заполните все поля\n"
        "3. Получите готовый документ!"
    )

def main():
    """
    Основная функция запуска бота
    """
    # Создаем приложение бота
    app = Application.builder().token("8299268017:AAHByyacpA819Q0HobJczPNmVrRIsGxPsog").build()
    
    # Настраиваем обработчик диалога
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler('start', start)],
        states={
            s_template: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_template_selection)],
            f_data: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_data_input)],
        },
        fallbacks=[CommandHandler('cancel', cancel)]
    )
    
    # Добавляем обработчики
    app.add_handler(conv_handler)
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("start", start))
    
    # Запускаем бота
    print("🤖 Бот для служебных записок запущен!")
    app.run_polling()

if __name__ == "__main__":
    main()